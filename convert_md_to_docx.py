import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE = os.path.join(BASE_DIR, "Vision_Series_Production_Test_Plan.md")
OUT_FILE = os.path.join(BASE_DIR, "Vision_Series_Production_Test_Plan.docx")
TEMPLATE_FILE = os.path.join(BASE_DIR, "templates", "word_document_template.docx")

COVER_TITLE = "High-level Production Test Plan"
COVER_SUBJECT = "Vision Series"

doc = Document(TEMPLATE_FILE)

# Strip template sample content but keep cover page AND front-matter (TOC, LoF, LoT)
body = doc.element.body
_cover_end = None
for _el in body:
    if _el.tag == qn('w:p'):
        _pPr = _el.find(qn('w:pPr'))
        if _pPr is not None and _pPr.find(qn('w:sectPr')) is not None:
            _cover_end = _el
            break

_to_remove = []
_past_cover = False
_remove_next_table = False
_removing_sample = False
for _el in list(body):
    if _el == _cover_end:
        _past_cover = True
        continue
    if not _past_cover:
        continue
    if _el.tag == qn('w:sectPr'):
        continue

    if _el.tag == qn('w:p'):
        _pPr = _el.find(qn('w:pPr'))
        _style = ''
        if _pPr is not None:
            _pStyle = _pPr.find(qn('w:pStyle'))
            if _pStyle is not None:
                _style = _pStyle.get(qn('w:val'), '')
        _text = ''.join(t.text or '' for t in _el.iter() if t.tag == qn('w:t'))
        if _style == 'Heading1-nonumber' and 'revision' in _text.lower():
            _to_remove.append(_el)
            _remove_next_table = True
            continue
        if _style == 'Heading1':
            _removing_sample = True

    if _remove_next_table and _el.tag == qn('w:tbl'):
        _to_remove.append(_el)
        _remove_next_table = False
        continue

    if _removing_sample:
        _to_remove.append(_el)

for _el in _to_remove:
    body.remove(_el)

# --- Fill in cover-page fields from the template ---
def _replace_cover_text(element, old, new):
    """Replace *old* with *new* inside every w:t run of *element*."""
    for t in element.iter(qn('w:t')):
        if t.text and old in t.text:
            t.text = t.text.replace(old, new)

def _get_latest_revision():
    """Parse the Revision History table and return (version, date) of the last row."""
    in_rev_table = False
    last = None
    with open(MD_FILE, "r", encoding="utf-8") as fh:
        for line in fh:
            if "Revision" in line and "Date" in line and "Description" in line:
                in_rev_table = True
                continue
            if in_rev_table:
                if line.startswith("|") and "---" not in line:
                    last = line
                elif not line.strip() or line.startswith("#"):
                    break
    cols = [c.strip() for c in last.strip().strip("|").split("|")]
    return cols[0], cols[1]

_rev_version, _rev_date = _get_latest_revision()

_cover_tbl = body[0]
if _cover_tbl.tag == qn('w:tbl'):
    _cover_rows = _cover_tbl.findall(qn('w:tr'))
    if len(_cover_rows) > 1:
        _replace_cover_text(_cover_rows[1], 'Project & (Type of document)', COVER_TITLE)
    if len(_cover_rows) > 2:
        _replace_cover_text(_cover_rows[2], '[Subject]', COVER_SUBJECT)
    # Row 3 has a nested version/date table; cells 1-3 are wrapped in w:sdt
    # content controls, so search the entire row for the placeholder text.
    if len(_cover_rows) > 3:
        _cell3 = _cover_rows[3].findall(qn('w:tc'))[0]
        _nested_tbls = _cell3.findall(qn('w:tbl'))
        if _nested_tbls:
            _nested_rows = _nested_tbls[0].findall(qn('w:tr'))
            if len(_nested_rows) > 1:
                _data_row = _nested_rows[1]
                _first_tc = _data_row.find(qn('w:tc'))
                if _first_tc is not None:
                    _replace_cover_text(_first_tc, '...', _rev_version)
                _replace_cover_text(_data_row, '[Publish Date]', _rev_date)

with open(MD_FILE, "r", encoding="utf-8") as f:
    lines = f.readlines()

def add_image(path_rel, caption=None):
    img_path = os.path.join(BASE_DIR, path_rel.replace("/", os.sep))
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(img_path, width=Inches(5.5))
        except Exception:
            run.add_picture(img_path, width=Inches(4.5))
        if caption:
            add_figure_caption(caption)
    else:
        p = doc.add_paragraph(f"[Image not found: {path_rel}]")
        p.runs[0].font.color.rgb = RGBColor(0xff, 0x00, 0x00)

STAGE_COL_NAMES = {"FT", "TRIM", "OT", "VIS", "VI", "VER", "CONN"}
STAGE_COL_WIDTH = Cm(1.1)
TEST_PROC_COL_WIDTH = Cm(3.0)
CONDITIONS_COL_WIDTH = Cm(2.5)
PAGE_CONTENT_WIDTH = Cm(17)  # 21cm A4 minus 2 × 2.0cm template margins

SPEC_TABLE_WIDTHS = {
    "Name": Cm(3.5),
    "Unit": Cm(1.5),
    "LSL": Cm(1.5),
    "USL": Cm(1.5),
}

def _add_seq_caption(style_name, seq_identifier, label, caption_text):
    """Add a caption paragraph using the template's Caption style with SEQ auto-numbering."""
    p = doc.add_paragraph(style=style_name)

    p.add_run(label + " ")

    r_begin = OxmlElement('w:r')
    fc_begin = OxmlElement('w:fldChar')
    fc_begin.set(qn('w:fldCharType'), 'begin')
    r_begin.append(fc_begin)
    p._p.append(r_begin)

    r_instr = OxmlElement('w:r')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' SEQ {} \\* ARABIC '.format(seq_identifier)
    r_instr.append(instr)
    p._p.append(r_instr)

    r_sep = OxmlElement('w:r')
    fc_sep = OxmlElement('w:fldChar')
    fc_sep.set(qn('w:fldCharType'), 'separate')
    r_sep.append(fc_sep)
    p._p.append(r_sep)

    r_num = OxmlElement('w:r')
    t_num = OxmlElement('w:t')
    t_num.text = '#'
    r_num.append(t_num)
    p._p.append(r_num)

    r_end = OxmlElement('w:r')
    fc_end = OxmlElement('w:fldChar')
    fc_end.set(qn('w:fldCharType'), 'end')
    r_end.append(fc_end)
    p._p.append(r_end)

    p.add_run(": " + caption_text)


def add_table_caption(caption_text):
    """Add a Word table caption with SEQ auto-numbering using the template Caption style."""
    _add_seq_caption('Caption', 'Table', 'Table', caption_text)


def add_figure_caption(caption_text):
    """Add a Word figure caption with SEQ auto-numbering using the template Caption style."""
    m = re.match(r'^Figure\s+\d+:\s*(.+)$', caption_text)
    if m:
        caption_text = m.group(1)
    _add_seq_caption('Caption', 'Figure', 'Figure', caption_text)

def add_table(header_line, rows):
    cols = [c.strip() for c in header_line.strip().strip("|").split("|")]
    clean_cols = [c.replace("**", "") for c in cols]
    num_cols = len(cols)
    table = doc.add_table(rows=1, cols=num_cols, style='Grid Table 4 Accent 3')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    is_test_proc_table = "Description" in clean_cols and "Test procedure" in clean_cols
    is_spec_table = "Name" in clean_cols and "Description" in clean_cols and "LSL" in clean_cols
    col_widths = None
    stage_col_indices = set()

    if is_spec_table:
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        layout_el = OxmlElement('w:tblLayout')
        layout_el.set(qn('w:type'), 'fixed')
        tbl_pr.append(layout_el)

        col_widths = []
        fixed_total = Cm(0)
        flex_count = 0
        for ci, name in enumerate(clean_cols):
            if name in SPEC_TABLE_WIDTHS:
                w = SPEC_TABLE_WIDTHS[name]
                col_widths.append(w)
                fixed_total += w
            else:
                col_widths.append(None)
                flex_count += 1
        remaining = PAGE_CONTENT_WIDTH - fixed_total
        if flex_count > 0:
            each = int(remaining) // flex_count
            col_widths = [w if w is not None else each for w in col_widths]

    elif is_test_proc_table:
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        layout_el = OxmlElement('w:tblLayout')
        layout_el.set(qn('w:type'), 'fixed')
        tbl_pr.append(layout_el)

        col_widths = []
        fixed_total = Cm(0)
        flex_count = 0
        for ci, name in enumerate(clean_cols):
            if name in STAGE_COL_NAMES:
                col_widths.append(STAGE_COL_WIDTH)
                fixed_total += STAGE_COL_WIDTH
                stage_col_indices.add(ci)
            elif name == "Test procedure":
                col_widths.append(TEST_PROC_COL_WIDTH)
                fixed_total += TEST_PROC_COL_WIDTH
            elif name == "Conditions":
                col_widths.append(CONDITIONS_COL_WIDTH)
                fixed_total += CONDITIONS_COL_WIDTH
            else:
                col_widths.append(None)
                flex_count += 1
        remaining = PAGE_CONTENT_WIDTH - fixed_total
        if flex_count > 0:
            each = int(remaining) // flex_count
            col_widths = [w if w is not None else each for w in col_widths]

    for i, col in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(col.replace("**", ""))
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        if col_widths:
            cell.width = col_widths[i]
        if i in stage_col_indices:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for row_text in rows:
        cells_text = [c.strip() for c in row_text.strip().strip("|").split("|")]
        row = table.add_row()

        # Detect section header rows: first cell bold, all others empty
        first_raw = cells_text[0] if cells_text else ""
        is_section_header = (
            first_raw.startswith("**") and first_raw.endswith("**")
            and all(c.strip() == "" for c in cells_text[1:num_cols])
        )

        if is_section_header and num_cols > 1:
            merged = row.cells[0].merge(row.cells[num_cols - 1])
            merged.text = ""
            run = merged.paragraphs[0].add_run(first_raw.replace("**", ""))
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D9E2F3')
            shading.set(qn('w:val'), 'clear')
            merged.paragraphs[0].runs[0]  # ensure paragraph exists
            tc_pr = merged._tc.get_or_add_tcPr()
            tc_pr.append(shading)
        else:
            for i, ct in enumerate(cells_text):
                if i < num_cols:
                    clean = ct.replace("**", "")
                    row.cells[i].text = clean
                    if col_widths:
                        row.cells[i].width = col_widths[i]
                    for para in row.cells[i].paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)
                            run.font.name = 'Calibri'
                    if i in stage_col_indices:
                        for para in row.cells[i].paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        for para in row.cells[i].paragraphs:
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

def parse_inline(text):
    """Returns list of (text, bold, italic) tuples."""
    parts = []
    pattern = re.compile(r'\*\*(.+?)\*\*')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False, False))
        parts.append((m.group(1), True, False))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False, False))
    return parts

i = 0
table_header = None
table_rows = []
in_table = False
pending_table_caption = None

while i < len(lines):
    line = lines[i].rstrip("\n")

    if line.startswith("|") and "---" not in line:
        if not in_table:
            in_table = True
            table_header = line
            table_rows = []
            i += 1
            if i < len(lines) and "---" in lines[i]:
                i += 1
            continue
        else:
            table_rows.append(line)
            i += 1
            continue
    else:
        if in_table:
            if pending_table_caption:
                add_table_caption(pending_table_caption)
                pending_table_caption = None
            add_table(table_header, table_rows)
            in_table = False
            table_header = None
            table_rows = []

    if line.strip() == "---":
        doc.add_paragraph()
        i += 1
        continue

    if line.startswith("# ") and not line.startswith("## "):
        doc.add_heading(line[2:].strip(), level=0)
        i += 1
        continue

    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
        i += 1
        continue

    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=2)
        i += 1
        continue

    if line.startswith("##### "):
        doc.add_heading(line[6:].strip(), level=4)
        i += 1
        continue

    if line.startswith("#### "):
        doc.add_heading(line[5:].strip(), level=3)
        i += 1
        continue

    img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
    if img_match:
        alt_text = img_match.group(1)
        img_path = img_match.group(2)
        caption = None
        if i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            cap_match = re.match(r'^\*(.+)\*$', next_line)
            if cap_match:
                caption = cap_match.group(1)
                i += 1
        add_image(img_path, caption)
        i += 1
        continue

    table_cap_match = re.match(r'^\*Table:\s*(.+)\*$', line.strip())
    if table_cap_match:
        pending_table_caption = table_cap_match.group(1).strip()
        i += 1
        continue

    cap_match = re.match(r'^\*([^*].+?[^*])\*$', line.strip())
    if cap_match and not line.strip().startswith("*Production") and not line.strip().startswith("*Trim") and not line.strip().startswith("*Stand") and not line.strip().startswith("*VIS Sub") and not line.strip().startswith("*(Similar") and not line.strip().startswith("*PSU") and not line.strip().startswith("*VIS4"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cap_match.group(1))
        run.italic = True
        run.font.size = Pt(9)
        i += 1
        continue

    if line.startswith("- **") or line.startswith("- *"):
        content = line[2:].strip()
        p = doc.add_paragraph(style='List Paragraph')
        for text, bold, italic in parse_inline(content):
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        i += 1
        continue

    if line.startswith("- "):
        content = line[2:].strip()
        p = doc.add_paragraph(style='List Paragraph')
        for text, bold, italic in parse_inline(content):
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        i += 1
        continue

    stripped = line.strip()
    if (stripped.startswith("*") and stripped.endswith("*")
            and not stripped.startswith("**")):
        text = stripped[1:-1]
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        i += 1
        continue

    if line.strip():
        p = doc.add_paragraph()
        for text, bold, italic in parse_inline(line.strip()):
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
        i += 1
        continue

    i += 1

if in_table:
    if pending_table_caption:
        add_table_caption(pending_table_caption)
        pending_table_caption = None
    add_table(table_header, table_rows)

doc.save(OUT_FILE)
print(f"Word document saved to: {OUT_FILE}")
